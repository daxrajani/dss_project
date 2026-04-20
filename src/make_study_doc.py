#!/usr/bin/env python3
"""
Generate DSS Project Study Guide — Word Document (.docx)
COEN 6731 Distributed Software Systems · Concordia University · Winter 2026
Dax Rajani (40294118) · Harsh Ahuja (40301748)

Usage:
    source dss_env/bin/activate
    python src/make_study_doc.py
Output:
    DSS_Study_Guide.docx
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
except ImportError:
    print("ERROR: pip install python-docx"); sys.exit(1)

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
OUT_PATH    = os.path.join(PROJECT_DIR, "DSS_Study_Guide.docx")

# ── Colors (RGB) ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
BLUE   = RGBColor(0x1A, 0x5E, 0xA8)
TEAL   = RGBColor(0x0E, 0x7A, 0x8A)
GREEN  = RGBColor(0x1E, 0x6B, 0x3A)
ORANGE = RGBColor(0xB3, 0x55, 0x00)
RED    = RGBColor(0x9B, 0x1C, 0x1C)
GOLD   = RGBColor(0xB3, 0x76, 0x00)
GRAY   = RGBColor(0x44, 0x44, 0x44)
LGRAY  = RGBColor(0x88, 0x88, 0x88)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def new_doc() -> Document:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
    # Remove default empty paragraph
    for p in doc.paragraphs:
        p.clear()
    return doc


def set_font(run, name="Calibri", size=11, bold=False, italic=False,
             color=None, underline=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1, color=NAVY):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    set_font(run, size=sizes.get(level, 12), bold=True, color=color)
    if level == 1:
        # Bottom border for H1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A5EA8")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def para(doc, text="", bold=False, italic=False, size=11, color=BLACK,
         space_before=2, space_after=4, indent=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def bullet(doc, text, level=0, bold_prefix=None, size=11):
    """Add a bullet point. bold_prefix is bolded, rest is normal."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=size, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_font(r2, size=size, color=BLACK)
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=BLACK)
    return p


def code_block(doc, code_text, caption=None):
    """Add a monospace code block with light gray background."""
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after  = Pt(1)
        cr = cp.add_run(caption)
        set_font(cr, size=9, bold=True, italic=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)

    # Light gray shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F4F8")
    pPr.append(shd)

    run = p.add_run(code_text)
    set_font(run, name="Courier New", size=9, color=RGBColor(0x1A, 0x1A, 0x2E))
    return p


def info_box(doc, title, body_lines, color_hex="1A5EA8"):
    """Styled single-cell table acting as a callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)

    # Background color
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "E8F0FA")
    tcPr.append(shd)

    # Left border accent
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:color"), color_hex)
    tcBorders.append(left)
    tcPr.append(tcBorders)

    cell.paragraphs[0].clear()
    tp = cell.paragraphs[0]
    tr = tp.add_run(title + "  ")
    set_font(tr, size=11, bold=True, color=BLUE)

    for line in body_lines:
        bp = cell.add_paragraph(line)
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after  = Pt(1)
        for run in bp.runs:
            set_font(run, size=10.5, color=BLACK)

    doc.add_paragraph()  # spacing after box
    return tbl


def key_result_box(doc, items):
    """2-column table for Key Result highlights."""
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(items):
        lc = tbl.cell(i, 0)
        vc = tbl.cell(i, 1)
        # shade label cell
        set_cell_bg(lc, "0D1B3E")
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        set_font(lr, size=10, bold=True, color=WHITE)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        set_font(vr, size=10, color=BLACK)
    doc.add_paragraph()


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def data_table(doc, headers, rows, col_widths=None):
    """Generic bordered table with navy header row."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        set_cell_bg(cell, "1A5EA8")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row_data in enumerate(rows):
        drow = tbl.rows[i + 1]
        bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        for j, cell_val in enumerate(row_data):
            cell = drow.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_val))
            set_font(r, size=10, color=BLACK)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])

    doc.add_paragraph()
    return tbl


def section_break(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 90)
    set_font(run, size=8, color=LGRAY)


def subheading(doc, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=color)
    return p


def mixed_para(doc, parts, space_before=2, space_after=4, indent=None):
    """
    parts = list of (text, bold, color) tuples.
    Renders them all in one paragraph.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold, color in parts:
        r = p.add_run(text)
        set_font(r, size=11, bold=bold, color=color)
    return p


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT SECTIONS
# ─────────────────────────────────────────────────────────────────────────────

def cover_page(doc):
    para(doc, space_before=60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DSS Project — Complete Study Guide")
    set_font(r, size=26, bold=True, color=NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Scalable E-Commerce Analytics Pipeline with Apache Spark")
    set_font(r2, size=15, italic=True, color=BLUE)

    para(doc, space_before=20)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("COEN 6731 — Distributed Software Systems\nConcordia University, Winter 2026")
    set_font(r3, size=13, color=GRAY)

    para(doc, space_before=8)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Dax Rajani (40294118)  ·  Harsh Ahuja (40301748)")
    set_font(r4, size=12, bold=True, color=NAVY)

    para(doc, space_before=30)
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(
        "This document covers all project theory, implementation details,\n"
        "benchmark results, optimization techniques, and probable professor questions.\n"
        "Intended for deep study — not for submission."
    )
    set_font(r5, size=11, italic=True, color=GRAY)

    doc.add_page_break()


def section_overview(doc):
    heading(doc, "1.  Project Overview")

    para(doc,
         "This project implements a fully distributed e-commerce analytics pipeline "
         "using Apache Spark (PySpark) on Google Cloud Platform's Dataproc service. "
         "The pipeline processes the REES46 October 2019 dataset — 42+ million "
         "behavioral events from an online multi-category store — and extracts four "
         "categories of business insight: conversion funnels, marketing attribution, "
         "anomaly detection, and scaling behavior.", size=11)

    subheading(doc, "1.1  Problem Motivation")
    para(doc,
         "Traditional single-machine tools like pandas load all data into one machine's "
         "RAM. At 5.3 GB CSV and 42 million rows, pandas would require 15–20 GB RAM "
         "for the enriched dataset, and window operations on 1.4 million users would "
         "run single-threaded for hours. A distributed approach is required.", size=11)

    subheading(doc, "1.2  Pipeline at a Glance")
    info_box(doc, "5-Stage Pipeline",
             ["Stage 1 — ETL & Enrichment:  Read raw CSV, synthesize device/referrer, derive orders/catalog, write Parquet",
              "Stage 2 — Sessionization:     Group events into 30-min inactivity sessions per user using window functions",
              "Stage 3 — Funnel Analysis:    Compute view→cart→purchase conversion rates by category/device/referrer",
              "Stage 4 — Last-Touch Attribution:  For each order, credit the most recent non-direct referrer within 24 hours",
              "Stage 5 — Anomaly Detection:  7-day rolling z-score baseline; flag SPIKE/DROP days per-system and per-category"])

    subheading(doc, "1.3  Technology Choices")
    data_table(doc,
               ["Technology", "Choice", "Reason"],
               [
                   ["Processing Engine", "Apache Spark (PySpark)", "Distributed in-memory, SQL-like DataFrame API, runs on clusters"],
                   ["Intermediate Format", "Apache Parquet", "Columnar, compressed, partition-aware, schema-preserving"],
                   ["Cloud Platform", "GCP Dataproc", "Managed Spark clusters, easy resize, integrates with GCS"],
                   ["Object Storage", "Google Cloud Storage (GCS)", "Decoupled storage from compute; persist across cluster restarts"],
                   ["Language", "Python 3 (PySpark)", "PySpark bindings; easy to prototype and extend"],
               ])


def section_theory(doc):
    heading(doc, "2.  Theoretical Background & Core Concepts")

    subheading(doc, "2.1  Distributed Computing Fundamentals")
    para(doc,
         "Distributed computing breaks a large computation into smaller tasks that run "
         "concurrently across multiple machines (nodes). The key challenges are:", size=11)
    bullet(doc, "Data Partitioning: how to divide the dataset so each node works on a roughly equal share")
    bullet(doc, "Fault Tolerance: if one node fails mid-computation, results must not be lost")
    bullet(doc, "Data Locality: move computation to where the data lives, not the other way around")
    bullet(doc, "Shuffle Cost: when tasks on different nodes need each other's data, they must exchange it over the network — this is called a shuffle, and it is the most expensive operation in distributed processing")

    subheading(doc, "2.2  Apache Spark Architecture")
    para(doc,
         "Spark uses a master-worker architecture. The driver program (master) creates a "
         "SparkContext, builds an execution plan, and distributes tasks to executors "
         "(workers). Each executor runs multiple tasks in parallel on its CPU cores.", size=11)

    para(doc, "Key Spark components:", bold=True, size=11)
    bullet(doc, "Driver:  Runs the main program. Hosts SparkContext and the DAG Scheduler.")
    bullet(doc, "Executor: JVM process on each worker node. Holds data partitions in memory/disk and executes tasks.")
    bullet(doc, "DAG Scheduler: Converts the logical plan into physical stages separated by shuffle boundaries.")
    bullet(doc, "Task Scheduler: Assigns individual tasks to executors, respecting data locality.")
    bullet(doc, "Catalyst Optimizer: Rewrites the logical plan into an optimized physical plan (pushes filters down, eliminates unused columns, chooses join strategies).")

    code_block(doc,
               "SparkSession.builder\n"
               "  .appName('DSS_Pipeline')\n"
               "  .master('local[*]')        # local mode: all cores on one machine\n"
               "  # On Dataproc, .master() is set by YARN automatically\n"
               "  .config('spark.driver.memory', '6g')\n"
               "  .config('spark.sql.shuffle.partitions', '200')\n"
               "  .getOrCreate()",
               caption="SparkSession creation (benchmark_job.py)")

    subheading(doc, "2.3  RDDs vs DataFrames vs Datasets")
    para(doc,
         "Spark has three APIs. We use DataFrames throughout this project:", size=11)

    data_table(doc,
               ["API", "Description", "Used In Project?"],
               [
                   ["RDD (Resilient Distributed Dataset)", "Low-level, type-safe, no query optimizer", "No — too verbose"],
                   ["DataFrame", "Tabular, untyped Python API, full Catalyst optimization", "YES — all stages"],
                   ["Dataset", "Type-safe DataFrame, Java/Scala only", "No — Python-only project"],
               ])

    para(doc,
         "DataFrames are the right choice here because Catalyst can optimize across the "
         "entire query plan. For example, it pushes filter predicates before joins "
         "(predicate pushdown) and reads only needed columns from Parquet (column pruning).", size=11)

    subheading(doc, "2.4  Lazy Evaluation and the DAG")
    para(doc,
         "Spark uses lazy evaluation: transformations (filter, groupBy, join) do NOT "
         "execute immediately. They build a Directed Acyclic Graph (DAG) of the "
         "computation. Execution only happens when an action is called.", size=11)

    code_block(doc,
               "# These lines build the DAG — nothing runs yet:\n"
               "df2 = df.filter(col('event_type') == 'purchase')\n"
               "df3 = df2.groupBy('user_id').agg(count('*').alias('purchases'))\n"
               "df4 = df3.orderBy('purchases', ascending=False)\n\n"
               "# This ACTION triggers execution of the full DAG:\n"
               "df4.show(10)",
               caption="Lazy evaluation — actions trigger DAG execution")

    para(doc,
         "A DAG is split into stages at shuffle boundaries. Within a stage, tasks are "
         "pipelined (no data movement). Between stages, data is exchanged over the "
         "network (shuffle). Minimizing shuffles is a core Spark optimization strategy.", size=11)

    subheading(doc, "2.5  Window Functions")
    para(doc,
         "A window function computes a value for each row based on a 'window' of "
         "surrounding rows without collapsing the DataFrame (unlike groupBy/agg). "
         "This is essential for sessionization and anomaly detection.", size=11)

    para(doc, "Three types of window frames:", bold=True, size=11)
    bullet(doc, "ROWS BETWEEN: frame defined by row count (e.g., 3 rows before current)")
    bullet(doc, "RANGE BETWEEN: frame defined by value range on the ORDER BY column (e.g., seconds within a time range)")
    bullet(doc, "UNBOUNDED PRECEDING / FOLLOWING: extend to the start or end of the partition")

    code_block(doc,
               "from pyspark.sql.window import Window\n"
               "from pyspark.sql.functions import lag, sum, avg, stddev\n\n"
               "# Partition by user, ordered by time — used for sessionization\n"
               "user_win = Window.partitionBy('user_id').orderBy('event_time')\n\n"
               "# lag() gets the previous row's value in the window\n"
               "df = df.withColumn('prev_time', lag('event_time').over(user_win))\n\n"
               "# Running cumulative sum within the window\n"
               "df = df.withColumn('session_idx', sum('is_new_session').over(user_win))\n\n"
               "# rangeBetween uses VALUE differences on the ORDER BY column\n"
               "# -604800 = -7 days in seconds, -86400 = -1 day\n"
               "rolling_win = Window.orderBy('date_ts').rangeBetween(-604800, -86400)\n"
               "df = df.withColumn('rolling_avg', avg('purchases').over(rolling_win))",
               caption="Window function examples — used in Stages 2 and 5")

    subheading(doc, "2.6  Apache Parquet Format")
    para(doc,
         "Parquet is a columnar binary file format designed for analytics workloads. "
         "Unlike row-based CSV (which stores complete rows together), Parquet stores "
         "each column's data contiguously.", size=11)

    para(doc, "Why Parquet outperforms CSV for analytics:", bold=True, size=11)
    bullet(doc, "Column Pruning: if a query only needs 3 of 15 columns, Parquet reads only those 3 — CSV always reads the full row")
    bullet(doc, "Compression: data within a single column is more homogeneous (same type, often similar values) → compresses far better. Our CSV is 5.3 GB, Parquet is ~1.2 GB")
    bullet(doc, "Predicate Pushdown: Parquet stores per-column min/max statistics in the footer; Spark can skip entire row groups without reading them")
    bullet(doc, "Schema Embedding: column names, types, and nullability are stored in the file — no need to infer types")
    bullet(doc, "Partition Pruning: Parquet files can be organized into directory partitions (e.g., date=2019-10-01/); Spark skips irrelevant directories entirely")

    code_block(doc,
               "# Write with date partitioning\n"
               "df.write.partitionBy('date').parquet('gs://bucket/output/')\n\n"
               "# Result on GCS:\n"
               "# output/date=2019-10-01/part-00000.parquet\n"
               "# output/date=2019-10-02/part-00000.parquet\n"
               "# ... (31 directories for October)\n\n"
               "# Query: Spark opens ONLY date=2019-10-15/ — skips other 30 dirs\n"
               "df.filter(col('date') == '2019-10-15').count()\n"
               "# → 3.11x speedup vs flat Parquet",
               caption="Parquet partition pruning — partition_analysis.py")

    subheading(doc, "2.7  Shuffle Join vs Broadcast Join")
    para(doc,
         "When Spark joins two DataFrames on a key, it must ensure that matching keys "
         "end up on the same executor. There are two main strategies:", size=11)

    para(doc, "SortMergeJoin (default — 'shuffle join'):", bold=True, size=11)
    bullet(doc, "Both tables are repartitioned by the join key (hash partitioning)")
    bullet(doc, "Spark sends rows with the same hash to the same executor across the network")
    bullet(doc, "Each executor then sorts and merges its local data")
    bullet(doc, "Cost: O(N log N) sort + network shuffle for BOTH tables")

    para(doc, "BroadcastHashJoin:", bold=True, size=11)
    bullet(doc, "The smaller table is fully sent (broadcast) to EVERY executor")
    bullet(doc, "Each executor builds a local hash map of the small table in memory")
    bullet(doc, "The larger table is scanned locally — no shuffle needed for it")
    bullet(doc, "Cost: broadcast small table once + local hash lookup. Only applicable when small table fits in executor memory")

    code_block(doc,
               "from pyspark.sql.functions import broadcast\n\n"
               "# Without broadcast: Spark uses SortMergeJoin (shuffles both sides)\n"
               "result = orders.join(session_refs, on='user_id')         # 11.33s\n\n"
               "# With broadcast hint: orders (153K rows) broadcast to all workers\n"
               "result = broadcast(orders).join(session_refs, on='user_id')  # 7.15s\n"
               "# → 1.58x speedup",
               caption="Broadcast join vs shuffle join — broadcast_join_opt.py")

    subheading(doc, "2.8  Data Skew and Salting")
    para(doc,
         "Data skew occurs when some partition keys are much more frequent than others, "
         "causing some Spark tasks to process far more data than others. The overloaded "
         "tasks become stragglers — the entire job waits for them to finish.", size=11)

    para(doc, "Salting technique:", bold=True, size=11)
    bullet(doc, "Append a random integer (salt) to the hot key to distribute it across N partitions")
    bullet(doc, "Replicate the smaller side of the join N times (one copy per salt value)")
    bullet(doc, "Join on (original_key, salt) instead of original_key alone")
    bullet(doc, "Result: the hot key's rows are split across N tasks instead of 1")

    code_block(doc,
               "N = 20  # salt buckets\n\n"
               "# Add random salt to orders\n"
               "salted_orders = orders.withColumn('salt', floor(rand() * N))\n"
               "salted_orders = salted_orders.withColumn(\n"
               "    'salted_key', concat(col('user_id'), lit('_'), col('salt')))\n\n"
               "# Replicate session_refs N times with matching salt values\n"
               "salted_sessions = session_refs.crossJoin(\n"
               "    spark.range(N).withColumnRenamed('id', 'salt'))\n"
               "salted_sessions = salted_sessions.withColumn(\n"
               "    'salted_key', concat(col('user_id'), lit('_'), col('salt')))\n\n"
               "# Join on salted key\n"
               "result = salted_orders.join(salted_sessions, on='salted_key')",
               caption="Salting implementation — skew_analysis.py")

    para(doc,
         "In our dataset, the top-1 user was only 0.02% of events — no meaningful "
         "skew existed. The crossJoin to replicate orders 20× added more overhead "
         "(79.1s) than the baseline without salting (34.6s). This confirmed a key "
         "lesson: always measure skew before applying mitigation.", size=11)

    subheading(doc, "2.9  Z-Score and Statistical Anomaly Detection")
    para(doc,
         "A z-score measures how many standard deviations a data point is from the mean "
         "of a reference distribution:", size=11)

    code_block(doc,
               "z = (observed_value - mean) / standard_deviation\n\n"
               "# z =  0.0 : exactly average\n"
               "# z =  2.0 : 2 standard deviations ABOVE average  → flag as SPIKE\n"
               "# z = -2.0 : 2 standard deviations BELOW average  → flag as DROP\n"
               "# |z| > 2 covers ~95% confidence interval — 5% of normal days flagged",
               caption="Z-score formula and interpretation")

    para(doc,
         "We use a 7-day TRAILING rolling window for the baseline — only past data "
         "contributes to today's mean and stddev. This is appropriate for time-series "
         "anomaly detection because it adapts to trends without leaking future information.", size=11)

    subheading(doc, "2.10  Last-Touch Attribution Model")
    para(doc,
         "Marketing attribution is the problem of crediting a conversion (purchase) to "
         "one or more touchpoints in the customer journey. Several models exist:", size=11)

    data_table(doc,
               ["Model", "Description", "Example"],
               [
                   ["First-Touch", "100% credit to first channel visited", "User found via Google → Google gets credit even if they came back 3 days later via email"],
                   ["Last-Touch", "100% credit to the last channel before purchase", "User's most recent non-direct session was Facebook → Facebook gets credit"],
                   ["Linear", "Equal credit across all touchpoints", "3 channels visited → each gets 33%"],
                   ["Time-Decay", "More recent channels get more credit", "Exponential weighting toward purchase time"],
                   ["Data-Driven", "ML-based, based on historical conversion patterns", "Requires large training data"],
               ])
    para(doc,
         "We implement Last-Touch because it is straightforward, interpretable, and "
         "widely used in e-commerce. The 24-hour lookback window limits attribution "
         "to sessions immediately preceding the order, filtering out sessions that "
         "are too old to be meaningfully credited.", size=11)

    subheading(doc, "2.11  Session-Based User Behavior Modeling")
    para(doc,
         "A web session represents a continuous period of activity by a single user. "
         "The 30-minute inactivity threshold is the industry standard (used by Google "
         "Analytics). Sessions are important because:", size=11)
    bullet(doc, "They group events into meaningful units of intent (a user looking for a phone vs a user browsing shoes)")
    bullet(doc, "Attribution is performed per-session (not per-event): the referrer of the session gets credit, not each individual click")
    bullet(doc, "Funnel analysis is also per-session: a session either contains a purchase or it does not")

    subheading(doc, "2.12  CRC32 Hashing for Deterministic Synthesis")
    para(doc,
         "CRC32 (Cyclic Redundancy Check, 32-bit) is a hash function that maps any "
         "string to a 32-bit integer. Its key property for our use is determinism: "
         "the same input always produces the same output, regardless of when or where "
         "it is computed.", size=11)

    code_block(doc,
               "# PySpark built-in CRC32 function\n"
               "from pyspark.sql.functions import crc32\n\n"
               "# Same user_session UUID always → same device assignment\n"
               "bucket = crc32(col('user_session')) % 100\n"
               "# bucket is deterministically in [0, 99]\n\n"
               "device = when(bucket < 55, 'mobile')   # 55% of UUID space\n"
               "       .when(bucket < 85, 'desktop')   # 30%\n"
               "       .otherwise('tablet')             # 15%",
               caption="CRC32 for deterministic device synthesis — etl_enrichment.py")

    para(doc,
         "This ensures reproducibility: running the pipeline twice on the same data "
         "produces the same device/referrer values. A random assignment (rand()) would "
         "change every run, making benchmark comparisons meaningless.", size=11)


def section_dataset(doc):
    heading(doc, "3.  Dataset Deep Dive")

    subheading(doc, "3.1  Source and Scale")
    para(doc,
         "The REES46 eCommerce Behavior Data is a public Kaggle dataset capturing "
         "all user interactions on a large multi-category online store during "
         "October 2019. It contains 42,448,764 events.", size=11)

    data_table(doc,
               ["Metric", "Value"],
               [
                   ["Total events", "42,448,764"],
                   ["Purchase events", "742,849  (~1.75% of all events)"],
                   ["Distinct users", "~1,407,580"],
                   ["Distinct products", "166,794"],
                   ["Distinct categories", "~3,600"],
                   ["Time period", "October 1–31, 2019  (31 days)"],
                   ["Raw CSV size", "5.3 GB"],
                   ["Parquet size (after ETL)", "~1.2 GB (4.4× compression)"],
               ])

    subheading(doc, "3.2  Raw Schema")
    data_table(doc,
               ["Column", "Raw Type", "Notes"],
               [
                   ["event_time", "string (UTC)", "e.g. '2019-10-01 00:00:00 UTC' — stripped and parsed to timestamp"],
                   ["event_type", "string", "One of: view, cart, purchase"],
                   ["product_id", "string", "Numeric product identifier"],
                   ["category_id", "string", "Large integer ID — rarely used"],
                   ["category_code", "string", "Human-readable e.g. electronics.smartphone — often null"],
                   ["brand", "string", "e.g. samsung — often null"],
                   ["price", "string → double", "Cast to double in ETL"],
                   ["user_id", "string", "Anonymous user identifier"],
                   ["user_session", "string (UUID)", "Browser session UUID — basis for CRC32 synthesis"],
               ])

    subheading(doc, "3.3  Missing / Synthesized Fields")
    para(doc,
         "The raw dataset lacks device type and referrer channel — two critical "
         "dimensions for marketing analytics. We synthesize them deterministically "
         "using CRC32 hashing on user_session.", size=11)

    data_table(doc,
               ["Field", "Method", "Distribution"],
               [
                   ["device", "crc32(user_session) % 100", "mobile 55%,  desktop 30%,  tablet 15%"],
                   ["referrer", "crc32(user_session) % 100", "direct 40%,  google 25%,  facebook 15%,  instagram 8%,  email 7%,  affiliate 5%"],
               ])

    para(doc,
         "Note: both device and referrer are derived from the same CRC32 hash of "
         "user_session. They use different threshold ranges so they are correlated "
         "(same session → same device AND same referrer) but independently distributed "
         "across the user population.", size=11)

    subheading(doc, "3.4  Event Type Distribution")
    para(doc,
         "The funnel shape is visible in the event type breakdown:", size=11)
    bullet(doc, "view:     ~40M events  — users browsing")
    bullet(doc, "cart:     ~1.5M events — users adding to cart")
    bullet(doc, "purchase: ~742K events — completed orders")
    para(doc,
         "This roughly 54:2:1 ratio (view:cart:purchase) represents a typical "
         "e-commerce conversion funnel — most browsing sessions do not convert.", size=11)


def section_stage1(doc):
    heading(doc, "4.  Stage 1 — ETL & Enrichment")
    para(doc, "File: src/etl_enrichment.py  (229 lines)", italic=True, color=GRAY, size=10)

    subheading(doc, "4.1  Responsibilities")
    bullet(doc, "Read the 42M-row raw CSV from local disk or GCS")
    bullet(doc, "Fix timestamp format (strip trailing ' UTC', cast to TimestampType)")
    bullet(doc, "Synthesize device column using CRC32(user_session) % 100")
    bullet(doc, "Synthesize referrer column using same CRC32 hash with different thresholds")
    bullet(doc, "Derive orders table from purchase events only")
    bullet(doc, "Derive catalog table of distinct products")
    bullet(doc, "Write three Parquet tables to output directory")

    subheading(doc, "4.2  Step-by-Step Code Walkthrough")

    para(doc, "Step 1 — Read CSV without type inference:", bold=True, size=11)
    code_block(doc,
               "# inferSchema=False avoids a full extra pass over 42M rows\n"
               "raw = spark.read.csv(input_path, header=True, inferSchema=False)\n"
               "\n"
               "# Strip ' UTC' suffix and parse to TimestampType\n"
               "raw = raw.withColumn('event_time',\n"
               "    to_timestamp(regexp_replace(col('event_time'), r' UTC$', '')))\n"
               "raw = raw.withColumn('price', col('price').cast('double'))")

    para(doc, "Step 2 — Synthesize device:", bold=True, size=11)
    code_block(doc,
               "def _add_device(df):\n"
               "    bucket = crc32(col('user_session')) % 100\n"
               "    return df.withColumn('device',\n"
               "        when(bucket < 55, 'mobile')    # 55%\n"
               "        .when(bucket < 85, 'desktop')  # 30%\n"
               "        .otherwise('tablet')            # 15%\n"
               "    )")

    para(doc, "Step 3 — Derive orders table:", bold=True, size=11)
    code_block(doc,
               "orders = (\n"
               "    enriched\n"
               "    .filter(col('event_type') == 'purchase')\n"
               "    .select(\n"
               "        monotonically_increasing_id().alias('order_id'),\n"
               "        # monotonically_increasing_id() is faster than row_number()\n"
               "        # because it requires no sort — IDs are unique but not sequential\n"
               "        col('user_id'),\n"
               "        col('event_time').alias('order_time'),\n"
               "        col('price').alias('total_amount'),\n"
               "        col('user_session'),\n"
               "    )\n"
               ")")

    para(doc, "Step 4 — Derive catalog table:", bold=True, size=11)
    code_block(doc,
               "catalog = (\n"
               "    enriched\n"
               "    .select('product_id', 'category_code', 'brand', 'price')\n"
               "    .dropDuplicates(['product_id'])  # one row per unique product\n"
               "    .fillna({'category_code': 'unknown', 'brand': 'unknown'})\n"
               ")")

    para(doc, "Step 5 — Write Parquet (the bottleneck stage):", bold=True, size=11)
    code_block(doc,
               "enriched.write.mode('overwrite').parquet(events_out)\n"
               "orders.write.mode('overwrite').parquet(orders_out)\n"
               "catalog.write.mode('overwrite').parquet(catalog_out)\n"
               "\n"
               "# On 5 workers, 10 GB: write alone takes 224 seconds\n"
               "# On 2 workers, 10 GB: write alone takes 477 seconds (35% of pipeline)")

    subheading(doc, "4.3  Outputs")
    data_table(doc,
               ["Table", "Path", "Rows (full data)", "Description"],
               [
                   ["enriched_events", "data/processed/enriched_events/", "42,448,764", "All events + device + referrer"],
                   ["orders", "data/processed/orders/", "742,849", "Purchase events only + order_id"],
                   ["catalog", "data/processed/catalog/", "166,794", "Distinct products"],
               ])

    subheading(doc, "4.4  Key Design Decisions")
    bullet(doc, "Why inferSchema=False?  A second full pass over 42M rows just to guess types — waste of ~2 min")
    bullet(doc, "Why monotonically_increasing_id() for order_id?  Avoids a costly sort; IDs are unique globally across partitions")
    bullet(doc, "Why write Parquet before passing to next stages?  Avoids recomputing enrichment for every downstream stage; Parquet persists the cache to disk, not just RAM")


def section_stage2(doc):
    heading(doc, "5.  Stage 2 — Sessionization")
    para(doc, "File: src/attribution.py → _derive_sessions()  (reused by multiple modules)", italic=True, color=GRAY, size=10)

    subheading(doc, "5.1  What is a Session?")
    para(doc,
         "A session is a contiguous period of browsing by a single user. The industry "
         "standard (Google Analytics default) is a 30-minute inactivity timeout: if a "
         "user has no events for 30+ minutes, the next event begins a new session.", size=11)

    info_box(doc, "Example:",
             ["User A events:",
              "  10:00  view product      ← Session 1 starts",
              "  10:12  view product",
              "  10:25  add to cart       ← Session 1 continues",
              "  (31 min gap)",
              "  10:56  view product      ← Session 2 starts (gap > 30 min)",
              "  11:02  purchase          ← Session 2 continues",
              "  (2 hr gap)",
              "  13:15  view product      ← Session 3 starts"])

    subheading(doc, "5.2  Algorithm (4 window operations)")
    code_block(doc,
               "SECONDS_30_MIN = 1800\n"
               "\n"
               "# Window: per user, ordered by time\n"
               "user_win = Window.partitionBy('user_id').orderBy('event_time')\n"
               "\n"
               "# Step 1: find previous event time\n"
               "df = df.withColumn('_prev_ts', lag('event_time').over(user_win))\n"
               "# For the FIRST event per user, _prev_ts is NULL\n"
               "\n"
               "# Step 2: compute gap in seconds\n"
               "df = df.withColumn('_gap',\n"
               "    unix_timestamp('event_time') - unix_timestamp('_prev_ts'))\n"
               "\n"
               "# Step 3: flag new session (gap > 30 min OR first event)\n"
               "df = df.withColumn('_is_new',\n"
               "    when(col('_gap') > SECONDS_30_MIN, 1)\n"
               "    .when(col('_prev_ts').isNull(), 1)   # first event ever\n"
               "    .otherwise(0))\n"
               "\n"
               "# Step 4: cumulative sum of flags = session index per user\n"
               "df = df.withColumn('_session_idx', sum('_is_new').over(user_win))\n"
               "df = df.withColumn('session_id',\n"
               "    concat_ws('_', col('user_id'), col('_session_idx')))\n"
               "# session_id = '520088904_3' means user 520088904's 3rd session",
               caption="_derive_sessions() — attribution.py:38-61")

    subheading(doc, "5.3  Why cumulative sum works")
    para(doc,
         "Each time _is_new = 1, the running sum increments. All events between two "
         "increments share the same sum value, which becomes the session index. This "
         "is an elegant pattern for grouping consecutive runs without a loop.", size=11)
    code_block(doc,
               "Event  gap(s)  _is_new  _session_idx  session_id\n"
               "10:00   NULL      1          1          user_123_1\n"
               "10:12    720      0          1          user_123_1\n"
               "10:25    780      0          1          user_123_1\n"
               "10:56   1860      1          2          user_123_2  ← new session\n"
               "11:02    360      0          2          user_123_2",
               caption="Cumulative sum sessionization trace")

    subheading(doc, "5.4  Why _derive_sessions is in attribution.py")
    para(doc,
         "Attribution also needs sessions (to identify which referrer was active during "
         "a session). By putting the function in attribution.py and importing it into "
         "benchmark_job.py, we avoid code duplication. The sessionization logic is "
         "defined exactly once.", size=11)


def section_stage3(doc):
    heading(doc, "6.  Stage 3 — Funnel Analysis")
    para(doc, "File: src/benchmark_job.py → run_funnel_analysis()  (~35 lines)", italic=True, color=GRAY, size=10)

    subheading(doc, "6.1  What is a Conversion Funnel?")
    para(doc,
         "A conversion funnel models the steps users take from initial interest to "
         "a completed purchase. We compute two transition rates:", size=11)
    bullet(doc, "view_to_cart_pct:  % of sessions with a view that also had a cart add")
    bullet(doc, "cart_to_buy_pct:   % of sessions with a cart add that also had a purchase")

    subheading(doc, "6.2  Implementation")
    code_block(doc,
               "# Step 1: Tag each event\n"
               "df = df.withColumn('is_view',     when(col('event_type')=='view',1).otherwise(0))\n"
               "df = df.withColumn('is_cart',     when(col('event_type')=='cart',1).otherwise(0))\n"
               "df = df.withColumn('is_purchase', when(col('event_type')=='purchase',1).otherwise(0))\n"
               "\n"
               "# Step 2: Collapse to SESSION level using max()\n"
               "# max(is_view) = 1 if the session had ANY view, else 0\n"
               "session_flags = df.groupBy(\n"
               "    'session_id', 'date', 'week', 'category_code', 'device', 'referrer'\n"
               ").agg(\n"
               "    max('is_view').alias('has_view'),\n"
               "    max('is_cart').alias('has_cart'),\n"
               "    max('is_purchase').alias('has_purchase'),\n"
               ").cache()  # reused for daily AND weekly output\n"
               "\n"
               "# Step 3: Aggregate to daily conversion rates\n"
               "daily = session_flags.groupBy('date', 'category_code', 'device', 'referrer')\n"
               "    .agg(\n"
               "        sum('has_view').alias('total_sessions_with_view'),\n"
               "        sum('has_cart').alias('total_sessions_with_cart'),\n"
               "        sum('has_purchase').alias('total_sessions_with_purchase'),\n"
               "    ).withColumn('view_to_cart_pct',\n"
               "        round(col('total_sessions_with_cart')/col('total_sessions_with_view')*100, 2)\n"
               "    ).withColumn('cart_to_buy_pct',\n"
               "        round(col('total_sessions_with_purchase')/col('total_sessions_with_cart')*100, 2)\n"
               "    )",
               caption="Funnel analysis — benchmark_job.py:59-124")

    subheading(doc, "6.3  Why max() for session flags?")
    para(doc,
         "We want a binary flag: 'did this session contain at least one purchase?' "
         "Using max(is_purchase) where is_purchase ∈ {0,1} achieves this — it returns "
         "1 if any row had a purchase, 0 otherwise. Using sum() would count total "
         "purchases per session, which is not what we need for funnel conversion rates.", size=11)

    subheading(doc, "6.4  Output scale")
    bullet(doc, "Daily funnel rows: 28,460  (1 GB run)")
    bullet(doc, "Weekly funnel rows: fewer (7 weeks × categories × device × referrer)")
    bullet(doc, "session_flags is .cache() because it is used twice — daily and weekly computation")


def section_stage4(doc):
    heading(doc, "7.  Stage 4 — Last-Touch Attribution")
    para(doc, "File: src/attribution.py  (285 lines)", italic=True, color=GRAY, size=10)

    subheading(doc, "7.1  Attribution Logic")
    para(doc,
         "For each purchase order, we find the most recent non-direct referrer session "
         "that started within 24 hours before the order time.", size=11)

    info_box(doc, "Attribution Rule",
             ["Given an order at time T by user U:",
              "  1. Find all sessions by user U where referrer != 'direct'",
              "  2. Keep only sessions where session_start is between (T - 24h) and T",
              "  3. Pick the session with the LATEST session_start (most recent touchpoint)",
              "  4. Credit that session's referrer channel with the order"])

    subheading(doc, "7.2  Step-by-Step Code")

    para(doc, "Step 1 — Build session-referrer table:", bold=True, size=11)
    code_block(doc,
               "# Each (user_id, session_id) → session_start + referrer + device\n"
               "# first(referrer) is correct because CRC32 hash ensures all events\n"
               "# in a user_session have the same referrer value\n"
               "session_refs = (\n"
               "    sessionized_df\n"
               "    .groupBy('user_id', 'session_id')\n"
               "    .agg(\n"
               "        min('event_time').alias('session_start'),\n"
               "        first('referrer').alias('referrer'),\n"
               "        first('device').alias('device'),\n"
               "    )\n"
               ")")

    para(doc, "Step 2 — Filter out direct sessions (no credit for direct):", bold=True, size=11)
    code_block(doc,
               "candidates = session_refs.filter(col('referrer') != 'direct')")

    para(doc, "Step 3 — Left join with time window filters:", bold=True, size=11)
    code_block(doc,
               "joined = (\n"
               "    orders.alias('o')\n"
               "    .join(candidates.alias('s'), on='user_id', how='left')\n"
               "    # session must have started BEFORE the order\n"
               "    .filter(col('s.session_start') <= col('o.order_time'))\n"
               "    # and within 24 hours before the order\n"
               "    .filter(\n"
               "        (unix_timestamp('o.order_time') - unix_timestamp('s.session_start'))\n"
               "        <= 86400\n"
               "    )\n"
               ")\n"
               "# LEFT join: orders with no qualifying session have s.* = NULL\n"
               "# (unattributed orders — excluded from summary but not dropped)")

    para(doc, "Step 4 — Pick the most recent session using row_number():", bold=True, size=11)
    code_block(doc,
               "rank_win = Window.partitionBy('o.order_id').orderBy(\n"
               "    col('s.session_start').desc()  # DESC → rank 1 = most recent\n"
               ")\n"
               "\n"
               "attributed = (\n"
               "    joined\n"
               "    .withColumn('_rank', row_number().over(rank_win))\n"
               "    .filter(col('_rank') == 1)  # keep only the most recent session\n"
               "    .select('o.order_id', 'o.user_id', 'o.order_time',\n"
               "            'o.total_amount', 's.referrer', 's.session_start', 's.device')\n"
               ")",
               caption="Full attribution pipeline — attribution.py:89-131")

    subheading(doc, "7.3  Key Numbers")
    data_table(doc,
               ["Metric", "1 GB Run", "10 GB Run (2w)"],
               [
                   ["Orders attributed", "100,852", "1,051,458"],
                   ["Attribution join time", "22.3s  (2w)", "122.0s  (2w)"],
                   ["Attribution join time", "14.2s  (5w)", "58.1s   (5w)"],
               ])


def section_stage5(doc):
    heading(doc, "8.  Stage 5 — Anomaly Detection")
    para(doc, "File: src/anomaly_detection.py  (279 lines)", italic=True, color=GRAY, size=10)

    subheading(doc, "8.1  Algorithm Overview")
    para(doc,
         "For each day, compute a rolling 7-day baseline (trailing, excluding today). "
         "Calculate the z-score. If |z| > 2, flag the day as SPIKE or DROP.", size=11)

    code_block(doc,
               "z_score = (purchase_count - rolling_avg) / rolling_stddev\n"
               "\n"
               "is_anomaly = abs(z_score) > 2.0\n"
               "\n"
               "explanation = SPIKE if z > 2 else DROP if z < -2 else None",
               caption="Core anomaly logic")

    subheading(doc, "8.2  Rolling Window Implementation")
    code_block(doc,
               "SECONDS_7D = 7 * 86400   # 604,800\n"
               "SECONDS_1D = 1 * 86400   #  86,400\n"
               "\n"
               "# date_ts is UNIX timestamp of the date (integer seconds)\n"
               "df = df.withColumn('date_ts', unix_timestamp(to_date(col('date'))))\n"
               "\n"
               "# rangeBetween works on VALUE differences on the ORDER BY column\n"
               "# This window: rows where date_ts is in [today-7days, today-1day]\n"
               "rolling_win = (\n"
               "    Window.orderBy('date_ts')\n"
               "    .rangeBetween(-SECONDS_7D, -SECONDS_1D)\n"
               ")\n"
               "\n"
               "df = df.withColumn('rolling_avg',    avg('purchase_count').over(rolling_win))\n"
               "df = df.withColumn('rolling_stddev', stddev('purchase_count').over(rolling_win))\n"
               "\n"
               "# Guard: if stddev is 0 or null (flat baseline), z_score is undefined\n"
               "df = df.withColumn('z_score',\n"
               "    when(col('rolling_stddev').isNotNull() & (col('rolling_stddev') > 0),\n"
               "        round((col('purchase_count') - col('rolling_avg'))\n"
               "              / col('rolling_stddev'), 3)\n"
               "    ).otherwise(None)\n"
               ")",
               caption="_apply_rolling_zscore() — anomaly_detection.py:35-102")

    subheading(doc, "8.3  Human-Readable Explanations")
    code_block(doc,
               "pct_above = round((purchase_count - rolling_avg) / rolling_avg * 100, 1)\n"
               "\n"
               "explanation = when(is_anomaly AND z > 0,\n"
               "    concat('SPIKE: ', pct_above, '% above 7-day baseline (',\n"
               "           purchase_count, ' vs expected ', rolling_avg, ')')\n"
               ").when(is_anomaly AND z < 0,\n"
               "    concat('DROP: ', pct_below, '% below 7-day baseline (...)')\n"
               ")\n"
               "\n"
               "# Example output:\n"
               "# 'SPIKE: 45.2% above 7-day baseline (1200 vs expected 826)'",
               caption="Explanation generation — anomaly_detection.py:76-100")

    subheading(doc, "8.4  Per-Category Anomalies")
    para(doc,
         "The same rolling z-score is computed separately per category_code by adding "
         "category_code to the window partition:", size=11)
    code_block(doc,
               "# System-wide: no partitioning\n"
               "rolling_win = Window.orderBy('date_ts').rangeBetween(-SECONDS_7D, -SECONDS_1D)\n"
               "\n"
               "# Per-category: partition by category_code\n"
               "rolling_win = (\n"
               "    Window.partitionBy('category_code')\n"
               "    .orderBy('date_ts')\n"
               "    .rangeBetween(-SECONDS_7D, -SECONDS_1D)\n"
               ")\n"
               "# Each category gets its own independent rolling baseline")

    subheading(doc, "8.5  Key Numbers")
    bullet(doc, "System-level anomalies: 8  (1 GB run, 31 days in October)")
    bullet(doc, "Per-category anomalies: 305  (many categories, each with own baseline)")
    bullet(doc, "Anomaly stage runtime: ~16s  (5 workers, 1 GB)")
    bullet(doc, "Why z_score can be NULL: first 7 days of the month have no prior baseline; also when stddev=0 (flat counts)")


def section_optimizations(doc):
    heading(doc, "9.  Performance Optimizations")

    subheading(doc, "9.1  Broadcast Join  (src/broadcast_join_opt.py)")
    para(doc, "Motivation:", bold=True, size=11)
    para(doc,
         "The attribution join is orders ⊲⊳ session_referrers on user_id. "
         "The default SortMergeJoin shuffles both tables across the network — very "
         "expensive when session_referrers is large (millions of sessions). "
         "But orders is only ~153K rows on a 25M-event sample — small enough to fit "
         "entirely in each executor's memory.", size=11)

    code_block(doc,
               "# CRITICAL: disable Spark's auto-broadcast for a fair baseline\n"
               "spark.config('spark.sql.autoBroadcastJoinThreshold', '-1')\n"
               "\n"
               "# Baseline: Spark chooses SortMergeJoin (shuffles both sides)\n"
               "result = orders.join(session_refs, on='user_id')\n"
               "t_baseline = 11.33s\n"
               "\n"
               "# Optimized: broadcast(orders) tells Spark to copy orders to all workers\n"
               "result = broadcast(orders).join(session_refs, on='user_id')\n"
               "t_broadcast = 7.15s\n"
               "\n"
               "speedup = 11.33 / 7.15 = 1.58x",
               caption="Broadcast join comparison — broadcast_join_opt.py:103-125")

    key_result_box(doc, [
        ("Shuffle Join (baseline)", "11.33 seconds"),
        ("Broadcast Join", "7.15 seconds"),
        ("Speedup", "1.58×"),
        ("Orders attributed", "153,628 (identical both ways)"),
    ])

    para(doc,
         "Important: results are identical — broadcast join is a physical plan "
         "optimization, not a logical change. The join output is the same rows.", size=11)

    subheading(doc, "9.2  Partition Pruning  (src/partition_analysis.py)")
    para(doc, "Motivation:", bold=True, size=11)
    para(doc,
         "Date-range queries on purchase data (e.g., 'show me purchases on Oct 15') "
         "are common in analytics. Without partitioning, Spark must scan the entire "
         "dataset. With date-partitioned Parquet, Spark skips all other dates.", size=11)

    code_block(doc,
               "# Write FLAT (no partitioning)\n"
               "raw.write.parquet(out_flat)                       # 22.9s write\n"
               "# Files: part-00001.parquet, part-00002.parquet, ...\n"
               "\n"
               "# Write DATE-PARTITIONED\n"
               "raw.write.partitionBy('date').parquet(out_part)   # 26.0s write (+1.1s overhead)\n"
               "# Files: date=2019-10-01/part-00000.parquet\n"
               "#        date=2019-10-02/part-00000.parquet  ...\n"
               "\n"
               "# QUERY: single-day filter\n"
               "df_flat.filter(col('date') == '2019-10-15').count()  # 4.25s (full scan)\n"
               "df_part.filter(col('date') == '2019-10-15').count()  # 1.37s (1/31 files)\n"
               "# Speedup: 3.11x\n"
               "\n"
               "# QUERY: 7-day range filter\n"
               "df_flat.filter((col('date')>='2019-10-10')&(col('date')<='2019-10-17')).count()\n"
               "# flat: 3.09s    partitioned: 1.78s    speedup: 1.73x",
               caption="Partition pruning experiment — partition_analysis.py:90-150")

    key_result_box(doc, [
        ("Single-day query (flat)", "4.25 seconds"),
        ("Single-day query (partitioned)", "1.37 seconds  →  3.11× speedup"),
        ("7-day range query (flat)", "3.09 seconds"),
        ("7-day range query (partitioned)", "1.78 seconds  →  1.73× speedup"),
        ("Write overhead", "+3.1 seconds  (amortized quickly by query savings)"),
    ])

    subheading(doc, "9.3  Skew Analysis  (src/skew_analysis.py)")
    para(doc, "Motivation:", bold=True, size=11)
    para(doc,
         "Hot-key skew — where one join key has orders of magnitude more data than "
         "others — is a classic Spark performance problem. We test whether our "
         "attribution join has a skew problem.", size=11)

    code_block(doc,
               "# Detect hot keys — most frequent users by event count\n"
               "user_counts = enriched.groupBy('user_id').agg(count('*').alias('events'))\n"
               "top_users = user_counts.orderBy('events', ascending=False)\n"
               "\n"
               "# Result:\n"
               "#   Top-1 user  : 1,487 events  (0.02% of all events)\n"
               "#   Top-10 users: 6,462 events  (0.08% of all events)\n"
               "# → LOW SKEW DATASET — no dominant hot key",
               caption="Hot-key detection — skew_analysis.py")

    code_block(doc,
               "# Salting: add random bucket to join key\n"
               "N_BUCKETS = 20\n"
               "salted_orders = orders.withColumn('salt', floor(rand() * N_BUCKETS))\n"
               "salted_orders = salted_orders.withColumn('salted_key',\n"
               "    concat(col('user_id'), lit('_'), col('salt')))\n"
               "\n"
               "# Replicate session_refs 20 times (crossJoin)\n"
               "salted_sessions = session_refs.crossJoin(\n"
               "    spark.range(N_BUCKETS).withColumnRenamed('id', 'salt'))\n"
               "salted_sessions = salted_sessions.withColumn('salted_key',\n"
               "    concat(col('user_id'), lit('_'), col('salt')))\n"
               "\n"
               "result = salted_orders.join(salted_sessions, on='salted_key')\n"
               "# Runtime: 79.1s  vs  34.6s without salting  →  0.44x speedup (SLOWER)",
               caption="Salting implementation — skew_analysis.py")

    key_result_box(doc, [
        ("Without salting (baseline)", "34.6 seconds"),
        ("With salting (20 buckets)", "79.1 seconds  →  0.44× (2.28× SLOWER)"),
        ("Top-1 user share", "0.02%  — no real skew"),
        ("Root cause", "crossJoin replicates orders 20× — overhead > benefit"),
        ("Lesson", "Measure skew BEFORE applying mitigation"),
    ])


def section_cloud(doc):
    heading(doc, "10.  Cloud Deployment — GCP Dataproc")

    subheading(doc, "10.1  Architecture")
    para(doc,
         "Google Cloud Dataproc is a managed Apache Spark/Hadoop service. We provision "
         "a cluster with one master node and 2/4/5 worker nodes. Data lives on Google "
         "Cloud Storage (GCS), decoupled from the cluster.", size=11)

    code_block(doc,
               "Cluster: dss-bench-2w-10gb-14824\n"
               "Region:  us-central1\n"
               "Master:  1 × n1-standard-4  (4 vCPU, 15 GB RAM)\n"
               "Workers: 2 / 4 / 5 × n1-standard-4  (4 vCPU, 15 GB RAM each)\n"
               "\n"
               "GCS bucket: gs://dss-project-dax/\n"
               "  data/full/2019-Oct.csv      ← input\n"
               "  scripts/cloud_pipeline.py   ← uploaded script\n"
               "  results/                    ← output Parquet tables")

    subheading(doc, "10.2  Why Cloud Pipeline is Self-Contained")
    para(doc,
         "On Dataproc, a PySpark job submission looks like this:", size=11)
    code_block(doc,
               "gcloud dataproc jobs submit pyspark gs://bucket/scripts/cloud_pipeline.py \\\n"
               "    --cluster=dss-bench-2w-10gb-14824 \\\n"
               "    --region=us-central1 \\\n"
               "    -- --input gs://bucket/data/full/2019-Oct.csv \\\n"
               "       --output gs://bucket/results/run1 \\\n"
               "       --sample-fraction 0.2")
    para(doc,
         "If cloud_pipeline.py imported etl_enrichment.py, we would need to also upload "
         "that file and all its dependencies as --py-files. To keep the job submission "
         "simple, cloud_pipeline.py (554 lines) inlines all 5 stages into one file.", size=11)

    subheading(doc, "10.3  Sample Fraction Mapping")
    data_table(doc,
               ["Data Label", "Sample Fraction", "Approx Rows", "Approx Size"],
               [
                   ["1 GB",  "0.20  (20%)",  "~8.5M rows",  "~1 GB Parquet"],
                   ["5 GB",  "0.50  (50%)",  "~42.5M rows", "~2.6 GB Parquet"],
                   ["10 GB", "1.00  (100%)", "~84.9M rows", "~5.3 GB CSV input"],
               ])
    para(doc,
         "The labels (1GB/5GB/10GB) refer to approximate total data processed through "
         "the pipeline including intermediate Parquet files, not just the raw input.", size=11)

    subheading(doc, "10.4  Demo Script")
    code_block(doc,
               "bash demo.sh cloud    # Submit live pipeline job to GCP Dataproc\n"
               "bash demo.sh stream   # Real-time streaming anomaly demo (local)\n"
               "bash demo.sh results  # Print benchmark table + optimization results\n"
               "bash demo.sh all      # Cloud + results (default)")


def section_benchmarks(doc):
    heading(doc, "11.  Benchmark Results & Analysis")

    subheading(doc, "11.1  Full Benchmark Table")
    data_table(doc,
               ["Workers", "Data", "Total (s)", "ETL Write (s)", "Attr. Join (s)", "Rows/sec"],
               [
                   ["2", "1 GB",  "673.6",   "223.1", "22.3",  "12,607"],
                   ["2", "5 GB",  "1,021.8", "300.4", "44.1",  "41,543"],
                   ["2", "10 GB", "1,378.9", "477.2", "122.0", "61,569"],
                   ["4", "1 GB",  "282.4",   "101.4", "15.1",  "30,070"],
                   ["4", "5 GB",  "569.0",   "174.4", "54.4",  "74,602"],
                   ["4", "10 GB", "799.8",   "264.3", "67.3",  "106,148"],
                   ["5", "1 GB",  "256.9",   "88.7",  "14.2",  "33,055"],
                   ["5", "5 GB",  "409.7",   "123.0", "33.3",  "103,609"],
                   ["5", "10 GB", "674.0",   "223.9", "58.1",  "125,961"],
               ])

    subheading(doc, "11.2  Speedup Analysis (relative to 2-worker baseline)")
    data_table(doc,
               ["Dataset", "4-worker speedup", "5-worker speedup", "Ideal (2.5×)"],
               [
                   ["1 GB",  "2.38×", "2.62×", "2.50×  — 1GB exceeds ideal!"],
                   ["5 GB",  "1.80×", "2.49×", "2.50×  — close to ideal"],
                   ["10 GB", "1.72×", "2.05×", "2.50×  — I/O bottleneck shows"],
               ])

    subheading(doc, "11.3  ETL Write Dominates at Scale")
    para(doc,
         "On the 2-worker, 10 GB run, ETL Parquet write took 477 seconds out of "
         "1,379 seconds total — 35% of the entire pipeline. This is a key finding: "
         "Distributed I/O (not CPU) is the primary bottleneck at scale.", size=11)

    code_block(doc,
               "Stage          2w/10GB     5w/10GB     Reduction\n"
               "ETL write       477.2s      223.9s       2.13x\n"
               "Session write   413.8s      185.2s       2.23x\n"
               "Attr. join      122.0s       58.1s       2.10x\n"
               "TOTAL          1378.9s      674.0s       2.05x\n"
               "\n"
               "All stages scale sub-linearly because GCS bandwidth is shared",
               caption="Stage-level scaling — 2w to 5w on 10 GB")

    subheading(doc, "11.4  Throughput Scaling")
    data_table(doc,
               ["Configuration", "Throughput", "vs 2-worker 1GB baseline"],
               [
                   ["2 workers, 1 GB",  "12,607 rows/sec",  "1.00× (baseline)"],
                   ["2 workers, 10 GB", "61,569 rows/sec",  "4.88× (larger batches are efficient)"],
                   ["4 workers, 10 GB", "106,148 rows/sec", "8.42×"],
                   ["5 workers, 10 GB", "125,961 rows/sec", "9.99×"],
               ])
    para(doc,
         "Throughput increases with both cluster size AND data size. This is because "
         "Spark achieves better CPU utilization when there are more rows per partition "
         "to amortize task scheduling overhead.", size=11)

    subheading(doc, "11.5  Amdahl's Law and Sub-linear Speedup")
    para(doc,
         "Amdahl's Law states that the speedup from parallelizing a computation is "
         "limited by the sequential fraction of the work:", size=11)
    code_block(doc,
               "Speedup = 1 / (s + (1-s)/N)\n"
               "  where s = fraction of work that is sequential\n"
               "        N = number of parallel workers\n"
               "\n"
               "For 10 GB, ETL write is ~35% of the pipeline and is I/O-bound,\n"
               "not linearly parallelizable. Even with s=0.20 (20% sequential):\n"
               "  Speedup_max = 1 / (0.20 + 0.80/2.5) = 1 / (0.20 + 0.32) = 1.92x\n"
               "  Observed: 2.05x  (consistent — I/O is partially parallel)",
               caption="Amdahl's Law applied to our scaling results")


def section_qa(doc):
    heading(doc, "12.  Professor Questions — Deep-Dive Answers")

    qa_pairs = [
        (
            "Q1: Why use PySpark instead of pandas for this project?",
            "Pandas is single-machine and single-threaded. For 42M rows:\n"
            "  • Memory: enriched DataFrame requires ~15–20 GB RAM after adding device/referrer columns — exceeds a standard machine\n"
            "  • Window functions: sessionizing 1.4M users in pandas requires an O(n log n) sort per user group — hours on one core\n"
            "  • No parallelism: pandas uses one CPU core; Spark uses all cores on all workers\n\n"
            "Spark distributes both memory and computation. Each worker holds a partition of the data; window functions run in parallel across partitions."
        ),
        (
            "Q2: What is a shuffle and why is it expensive?",
            "A shuffle is the redistribution of data across nodes so that rows with the same key end up on the same executor. It happens during:\n"
            "  • groupBy / agg (rows with same group key must colocate)\n"
            "  • join (rows with same join key must colocate)\n"
            "  • orderBy / sort (all data must be globally sorted)\n\n"
            "Cost: each shuffle writes data to disk, serializes it, transmits over the network, deserializes, and reads from disk on the receiving end. "
            "This is 10–100× slower than in-memory operations. Minimizing shuffles is the primary Spark tuning goal."
        ),
        (
            "Q3: How does lazy evaluation help Spark optimize the query?",
            "When you write df.filter(...).groupBy(...).agg(...), Spark builds a logical plan (an AST). "
            "No data moves. The Catalyst optimizer then rewrites this plan:\n"
            "  • Predicate pushdown: filters are moved as early as possible (reduce rows before expensive operations)\n"
            "  • Column pruning: if downstream stages only use 3 of 15 columns, earlier stages only read those 3\n"
            "  • Join reordering: smaller tables are placed on the probe side\n"
            "  • Broadcast join selection: if one side is small, Catalyst automatically chooses BroadcastHashJoin\n\n"
            "Execution only begins when an action (.count(), .write(), .show()) is called, at which point the optimized physical plan runs."
        ),
        (
            "Q4: Why does the ETL stage take 35% of total runtime?",
            "ETL Parquet write time: 477s out of 1,379s total on 2 workers, 10 GB.\n\n"
            "Reasons:\n"
            "  1. The enriched_events table is the LARGEST (42M rows + 2 new columns). Writing Parquet requires compressing and serializing every row\n"
            "  2. GCS network bandwidth is shared across all workers — writing large amounts to cloud storage becomes I/O-bound\n"
            "  3. Parquet write is a shuffle-like operation: Spark repartitions the data before writing to ensure balanced file sizes\n\n"
            "This is a key lesson: at scale, I/O dominates, not computation."
        ),
        (
            "Q5: Why was salting slower in our dataset?",
            "Salting works by adding a random bucket integer (0..N-1) to the join key, distributing a hot key across N partitions. "
            "The technique requires:\n"
            "  1. crossJoin: replicate the smaller table N times (one copy per bucket)\n"
            "  2. Join on (user_id, salt) compound key instead of user_id alone\n\n"
            "In our dataset, the top-1 user has only 0.02% of events — no straggler problem exists. "
            "The crossJoin added 20× overhead to the orders table, and the compound key join produced no benefit. "
            "Result: 34.6s → 79.1s (0.44× speedup, i.e., 2.28× slower).\n\n"
            "Salting WOULD help if: top-1 user has >1–5% of events, causing one task to run 50–100× longer than others."
        ),
        (
            "Q6: What does rangeBetween(-604800, -86400) mean exactly?",
            "The rolling window is defined using rangeBetween on the date_ts column (UNIX timestamps in seconds).\n\n"
            "rangeBetween(A, B) includes all rows in the partition where:\n"
            "   current_row_date_ts + A  ≤  other_row_date_ts  ≤  current_row_date_ts + B\n\n"
            "So for today at ts = T:\n"
            "   Lower bound: T + (-604800) = T - 7 days\n"
            "   Upper bound: T + (-86400)  = T - 1 day\n\n"
            "Today (offset 0) is excluded from the baseline — this prevents today's anomalous value from inflating the rolling average and understating its own z-score."
        ),
        (
            "Q7: Why use left join in attribution instead of inner join?",
            "An inner join would DROP orders that have no qualifying referrer session. "
            "These are orders where the user only ever came directly (no paid/earned channel in the 24-hour window). "
            "Dropping them would be data loss — we want to know the total order count including unattributed orders.\n\n"
            "With left join, unattributed orders appear with s.* = NULL. "
            "They are excluded from the attribution summary (which only counts attributed channels) "
            "but remain in the attributed_orders table for completeness."
        ),
        (
            "Q8: Why does 1GB have super-linear speedup (2.62× with 2.5× ideal)?",
            "The 1 GB dataset has ~8.5M rows — relatively small for a 5-worker cluster. "
            "Each worker's tasks are small and complete quickly, meaning:\n"
            "  1. Better task scheduling efficiency — less time wasted on overheads\n"
            "  2. Data fits more comfortably in executor memory — fewer disk spills\n"
            "  3. The GCS I/O bottleneck is proportionally smaller (less data to write)\n\n"
            "Going from 2 to 5 workers on small data achieves near-linear gains because "
            "the sequential fraction (I/O bottleneck) is smaller. On large data (10 GB), "
            "the same bottleneck limits speedup to 2.05×."
        ),
        (
            "Q9: What is the difference between etl_enrichment.py and cloud_pipeline.py?",
            "etl_enrichment.py (229 lines):\n"
            "  • Does ONLY Stage 1 (ETL)\n"
            "  • Designed to be imported by benchmark_job.py\n"
            "  • Uses Python module imports for reuse\n\n"
            "cloud_pipeline.py (554 lines):\n"
            "  • Contains ALL 5 stages inlined in one file\n"
            "  • Self-contained — no local module imports\n"
            "  • Designed for gcloud dataproc jobs submit pyspark (single file submission)\n"
            "  • Accepts --sample-fraction argument for benchmarking different data sizes\n\n"
            "They implement the same logic. cloud_pipeline.py trades code reuse for deployment simplicity."
        ),
        (
            "Q10: How does Parquet partition pruning achieve 3.11× speedup on single-day queries?",
            "Without partitioning: Spark opens all Parquet files and reads every row, then filters.\n"
            "With date partitioning: Spark uses the directory structure as a filter index.\n\n"
            "For a single-day query on October (31 days), Spark:\n"
            "  • FLAT: opens and scans all 31 days of data  →  4.25s\n"
            "  • PARTITIONED: only opens date=2019-10-15/  →  1.37s\n\n"
            "The speedup is 31/1 × (I/O fraction of query time) ≈ 3× because I/O accounts "
            "for most of the query time. The remaining time (parsing, counting) is the same "
            "for both — that is why speedup is 3.11× and not 31×."
        ),
        (
            "Q11: What is Amdahl's Law and how does it explain our scaling results?",
            "Amdahl's Law: Speedup = 1 / (s + (1-s)/N)\n"
            "  s = fraction of work that cannot be parallelized\n"
            "  N = number of parallel processors\n\n"
            "In our pipeline, the ETL Parquet write is ~35% I/O-bound. I/O to GCS "
            "partially parallelizes (more workers = more parallel writes) but not perfectly. "
            "With s ≈ 0.18 (18% effectively sequential on 5-worker run):\n"
            "  Predicted speedup (2→5w) = 1/(0.18 + 0.82/2.5) = 1.92×\n"
            "  Observed: 2.05×  (reasonable agreement)\n\n"
            "This explains why adding the 5th worker gives diminishing returns vs adding the 3rd or 4th."
        ),
        (
            "Q12: What is the streaming component and how does it work?",
            "src/streaming_job.py uses Spark Structured Streaming to process events in real time.\n\n"
            "Architecture:\n"
            "  stream_producer.py  →  drops CSV batches to data/raw/stream_input/ every 2 sec\n"
            "  streaming_job.py    →  watches the directory as a streaming source\n"
            "                         applies 5-minute tumbling window aggregation\n"
            "                         flags SPIKE if purchase count > threshold\n"
            "                         injects a spike every 7th batch for demo effect\n\n"
            "Key Spark Streaming concept: micro-batch processing. Spark reads new files "
            "in the watch directory every trigger interval, processes them as a mini-batch, "
            "and outputs results. The checkpoint directory persists state between micro-batches."
        ),
    ]

    for q, a in qa_pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(q)
        set_font(r, size=12, bold=True, color=NAVY)

        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(1)
        ap.paragraph_format.space_after  = Pt(8)
        ap.paragraph_format.left_indent  = Inches(0.25)
        ar = ap.add_run(a)
        set_font(ar, size=11, color=BLACK)

        # Light divider
        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after  = Pt(0)
        dr = dp.add_run("─" * 80)
        set_font(dr, size=7, color=LGRAY)


def section_quick_ref(doc):
    heading(doc, "13.  Quick Reference — Numbers to Remember")

    subheading(doc, "Dataset")
    data_table(doc,
               ["Metric", "Value"],
               [
                   ["Total events", "42,448,764"],
                   ["Purchases", "742,849  (1.75%)"],
                   ["Distinct products", "166,794"],
                   ["Raw CSV", "5.3 GB"],
                   ["Date range", "October 2019  (31 days)"],
               ])

    subheading(doc, "Synthesis")
    data_table(doc,
               ["Field", "Distribution"],
               [
                   ["device",   "mobile 55%,  desktop 30%,  tablet 15%"],
                   ["referrer", "direct 40%,  google 25%,  facebook 15%,  instagram 8%,  email 7%,  affiliate 5%"],
               ])

    subheading(doc, "Algorithm Thresholds")
    data_table(doc,
               ["Parameter", "Value", "Why"],
               [
                   ["Session inactivity", "30 min  (1,800 s)", "Industry standard (Google Analytics default)"],
                   ["Attribution lookback", "24 hours  (86,400 s)", "Balance recency vs attribution window coverage"],
                   ["Anomaly z-threshold", "±2.0", "~95% confidence — flags statistically unusual days"],
                   ["Rolling baseline", "7 days (trailing)", "Captures weekly seasonality, excludes today"],
                   ["Salt buckets", "20", "Used in skew_analysis.py (proved unnecessary)"],
               ])

    subheading(doc, "Optimization Results")
    data_table(doc,
               ["Technique", "Before", "After", "Speedup"],
               [
                   ["Broadcast Join",             "11.33s", "7.15s",  "1.58×"],
                   ["Partition Pruning (1-day)",  "4.25s",  "1.37s",  "3.11×"],
                   ["Partition Pruning (7-day)",  "3.09s",  "1.78s",  "1.73×"],
                   ["Salting (skew mitigation)",  "34.6s",  "79.1s",  "0.44×  (SLOWER)"],
               ])

    subheading(doc, "Scaling Highlights")
    data_table(doc,
               ["Config", "Total Time", "Throughput"],
               [
                   ["2w / 1GB",  "673.6s",   "12,607 rows/sec"],
                   ["2w / 10GB", "1,378.9s", "61,569 rows/sec"],
                   ["4w / 10GB", "799.8s",   "106,148 rows/sec"],
                   ["5w / 10GB", "674.0s",   "125,961 rows/sec"],
                   ["Best speedup (2→5w)", "2.62× (1GB)", "82% of ideal 2.5×"],
               ])


# ─────────────────────────────────────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print("Building DSS Study Guide Word document ...")
    doc = new_doc()

    steps = [
        ("Cover page",            cover_page),
        ("1. Overview",           section_overview),
        ("2. Theory",             section_theory),
        ("3. Dataset",            section_dataset),
        ("4. Stage 1 ETL",        section_stage1),
        ("5. Stage 2 Sessions",   section_stage2),
        ("6. Stage 3 Funnel",     section_stage3),
        ("7. Stage 4 Attribution",section_stage4),
        ("8. Stage 5 Anomaly",    section_stage5),
        ("9. Optimizations",      section_optimizations),
        ("10. Cloud",             section_cloud),
        ("11. Benchmarks",        section_benchmarks),
        ("12. Q&A",               section_qa),
        ("13. Quick Reference",   section_quick_ref),
    ]

    for label, fn in steps:
        print(f"  Writing section: {label} ...", flush=True)
        fn(doc)

    doc.save(OUT_PATH)
    print(f"\nSaved → {OUT_PATH}")


if __name__ == "__main__":
    main()
